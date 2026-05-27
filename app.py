from dotenv import load_dotenv
load_dotenv()

import streamlit as st
import os
from PIL import Image
import fitz
from docx import Document
import io
import re
from openai import OpenAI
import base64

st.set_page_config(
    page_title="note2action — AI Meeting Notes Converter",
    page_icon="📋",
    layout="wide"
)

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=os.environ.get("OPENROUTER_API_KEY")
)

if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False

# CSS 
def get_css(dark):
    if dark:
        return """
        :root {
            --bg:        #0a0f0a;
            --bg2:       #111a11;
            --bg3:       transparent;
            --card-bg:   transparent;
            --border:    #2a3d2a;
            --text:      #e8f5e8;
            --text2:     #8fbc8f;
            --accent:    #4ade80;
            --accent2:   #86efac;
            --accent3:   #bbf7d0;
            --btn:       #166534;
            --btn-hover: #15803d;
            --success:   #14532d;
            --success-t: #86efac;
            --warn:      #713f12;
            --warn-t:    #fde68a;
            --metric-v:  #4ade80;
        }
        """
    else:
        return """
        :root {
            --bg:        #f0faf0;
            --bg2:       #e8f5e8;
            --bg3:       #dcf0dc;
            --border:    #a7d7a7;
            --text:      #1a2e1a;
            --text2:     #3d6b3d;
            --accent:    #16a34a;
            --accent2:   #15803d;
            --accent3:   #166534;
            --btn:       #16a34a;
            --btn-hover: #15803d;
            --card-bg:   #ffffff;
            --success:   #dcfce7;
            --success-t: #166534;
            --warn:      #fef9c3;
            --warn-t:    #713f12;
            --metric-v:  #16a34a;
        }
        """

st.markdown(f"""
<style>
{get_css(st.session_state.dark_mode)}

.block-container {{
    padding-top: 0.1rem !important;
    padding-bottom: 0rem !important;
    padding-left: 0.5rem !important;
    padding-right: 0.5rem !important;
    max-width: 100% !important;
}}

.main .block-container {{
    padding: 0 !important;
    margin: 0 !important;
}}

[data-testid="stAppViewContainer"] {{
    margin: 0 !important;
    padding: 0 !important;
}}

[data-testid="stAppViewContainer"] > .main {{
    padding: 0 !important;
    margin: 0 !important;
}}

section.main > div {{
    padding-top: 0 !important;
    padding-left: 0 !important;
    padding-right: 0 !important;
    max-width: 100% !important;
}}

.main-grid,
[data-testid="stHorizontalBlock"] {{
    margin-top: 0 !important;
    padding-top: 0 !important;
}}

.left-panel,
.right-panel {{
    padding-top: 10px !important;
}}

.element-container:first-child {{
    margin-top: 0 !important;
    padding-top: 0 !important;
}}

.stChatFloatingInputContainer {{ 
    display: none !important; 
}}
[data-testid="stAppViewContainer"] {{
    background: var(--bg) !important;
}}
[data-testid="stHeader"] {{
    background: var(--bg) !important;
    border-bottom: 1px solid var(--border);
}}
section[data-testid="stSidebar"] {{
    display: none;
}}

@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;500;600;700;800&display=swap');
html, body, [class*="css"] {{
    font-family: 'Plus Jakarta Sans', sans-serif;
    color: var(--text);
}}

.navbar {{
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 18px 48px;
    background: var(--bg2);
    border-bottom: 1px solid var(--border);
    position: sticky;
    top: 0;
    z-index: 100;
}}
.nav-logo {{
    font-size: 1.4rem;
    font-weight: 800;
    color: var(--accent);
    letter-spacing: -0.5px;
}}
.nav-logo span {{ color: var(--text); }}
.nav-badges {{ display: flex; gap: 8px; flex-wrap: wrap; }}
.badge {{
    padding: 4px 12px;
    border-radius: 20px;
    font-size: 0.72rem;
    font-weight: 600;
    background: var(--bg3);
    color: var(--accent2);
    border: 1px solid var(--border);
}}

/* Main layout */
.main-grid {{
    display: grid;
    grid-template-columns: 420px 1fr;
    gap: 0;
    min-height: calc(100vh - 65px);
}}

/* Left panel */
.left-panel {{
    background: var(--bg2);
    border-right: 1px solid var(--border);
    padding: 32px 28px;
}}
.panel-title {{
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: var(--text2);
    margin-bottom: 20px;
}}

/* Right panel */
.right-panel {{
    background: var(--bg);
    padding: 32px 36px;
}}

[data-testid="stTextInput"] input {{
    background: transparent !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 8px !important;
    color: var(--text) !important;
    font-family: 'Plus Jakarta Sans', sans-serif !important;
    padding: 10px 14px !important;
    font-size: 0.9rem !important;
    box-shadow: none !important;
}}
[data-testid="stTextInput"] input:focus {{
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 3px rgba(74,222,128,0.15) !important;
}}
[data-testid="stTextInput"] label {{
    color: var(--text2) !important;
    font-size: 0.8rem !important;
    font-weight: 600 !important;
}}
[data-testid="stTextInput"] input::placeholder {{
    color: var(--text2) !important;
    opacity: 1 !important;
}}
[data-testid="stTextInput"] div[data-baseweb="input"] {{
    background: transparent !important;
    background-color: transparent !important;
}}

[data-testid="stTextInput"] div[data-baseweb="base-input"] {{
    background: transparent !important;
    background-color: transparent !important;
}}

[data-testid="stTextInput"] > div {{
    background: transparent !important;
}}

[data-testid="stAlert"] div, [data-testid="stAlert"] p {{
    color: var(--text) !important;
    opacity: 1 !important;
}}
[data-testid="stTextInput"] {{
    margin-bottom: 0 !important;
    padding-bottom: 0 !important;
}}

[data-testid="stFileUploader"] {{
    background: var(--bg2) !important;
    border: 2px dashed var(--accent2) !important;
    border-radius: 16px !important;
    padding: 14px !important;
}}

[data-testid="stFileUploader"] section {{
    background: var(--bg3) !important;
    border-radius: 14px !important;
    border: 1px solid var(--border) !important;
}}

[data-testid="stFileUploader"] button,
[data-testid="stFileUploader"] button[kind],
[data-testid="stFileUploader"] .st-emotion-cache-1erivf3 button {{
    background: #15803d !important;
    background-color: #15803d !important;
    color: #ffffff !important;
    border: 1px solid #22c55e !important;
    border-radius: 16px !important;
    font-weight: 700 !important;
    font-size: 1rem !important;
    padding: 14px 30px !important;
    box-shadow: none !important;
    transition: all 0.2s ease !important;
}}

/* Hover */
[data-testid="stFileUploader"] button:hover {{
    background: #166534 !important;
    background-color: #166534 !important;
    color: #ffffff !important;
    border: 1px solid #4ade80 !important;
}}

[data-testid="stFileUploader"] button p {{
    color: #ffffff !important;
}}

[data-testid="stFileUploader"] button svg {{
    fill: #ffffff !important;
    stroke: #ffffff !important;
}}

[data-testid="stFileUploaderDropzoneInstructions"] * {{
    color: var(--text2) !important;
}}

[data-testid="stFileUploaderDropzone"] button {{
    background: #15803d !important;
    color: white !important;
    border: 1px solid #22c55e !important;
    border-radius: 16px !important;
}}

.stFileChip {{
    background: transparent !important;
    color: var(--text) !important;
    border: none !important;
}}

.stFileChip span,
.stFileChip small,
.stFileChip p {{
    color: var(--text) !important;
}}

.stFileChip svg {{
    color: var(--text) !important;
}}

.stFileChip [data-testid="stFileUploaderFile"] svg,
.stFileChip div:first-child {{
    background: transparent !important;
    border-radius: 10px !important;
    padding: 6px !important;
}}

.stFileChip,
.stFileChip * {{
    color: var(--text) !important;
}}

.stFileChip small,
.stFileChip span {{
    color: var(--text2) !important;
}}

.stFileChip div:first-child {{
    background: transparent !important;
    border-radius: 14px !important;
    padding: 8px !important;
}}

.stFileChip svg {{
    color: var(--text) !important;
    fill: none !important;
    stroke: var(--text) !important;
}}

[data-testid="stFileUploaderDropzone"] button p {{
    color: white !important;
}}

[data-testid="stFileUploaderDropzone"] button svg {{
    fill: white !important;
    stroke: white !important;
}}

[data-testid="stFileUploaderDropzone"] button p {{
    color: #ffffff !important;
}}

[data-testid="stFileUploader"] .st-emotion-cache-1umgz6k {{
    display: none !important;
}}

[data-testid="stFileUploader"] button svg {{
    fill: #ffffff !important;
    stroke: #ffffff !important;
}}

[data-testid="stFileUploader"] button[aria-label="Show help"] {{
    display: none !important;
}}

[data-testid="stFileUploaderDropzone"] {{
    background: var(--bg3) !important;
    border-radius: 14px !important;
}}

[data-testid="stFileUploaderFile"] {{
    background: var(--card-bg) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    color: var(--text) !important;
}}

[data-testid="stFileUploaderFile"] {{
    background: var(--bg2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 16px !important;
    overflow: hidden !important;
}}

[data-testid="stFileUploaderFile"] > div,
[data-testid="stFileUploaderFile"] section,
[data-testid="stFileUploaderFile"] article {{
    background: var(--bg2) !important;
}}

[data-testid="stFileUploaderFile"] * {{
    color: var(--text) !important;
}}

[data-testid="stFileUploaderFile"] button {{
    background: var(--btn) !important;
    border-radius: 12px !important;
}}


[data-testid="stButton"] button {{
    background: var(--btn) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
    padding: 12px 24px !important;
    width: 100% !important;
    transition: all 0.2s !important;
    font-family: 'Plus Jakarta Sans', sans-serif !important;
    letter-spacing: 0.3px !important;
}}
[data-testid="stButton"] button:hover {{
    background: var(--btn-hover) !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 20px rgba(74,222,128,0.25) !important;
}}

.metrics-row {{
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: 12px;
    margin-bottom: 24px;
}}
.metric-card {{
    background: var(--card-bg);
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 16px;
    text-align: center;
}}
.metric-value {{
    font-size: 1.8rem;
    font-weight: 800;
    color: var(--metric-v);
    line-height: 1;
}}
.metric-label {{
    font-size: 0.72rem;
    color: var(--text2);
    margin-top: 6px;
    font-weight: 500;
}}

[data-testid="stMarkdownContainer"] h1,
[data-testid="stMarkdownContainer"] h2,
[data-testid="stMarkdownContainer"] h3,
[data-testid="stMarkdownContainer"] h4,
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stMarkdownContainer"] strong {{
    color: var(--text) !important;
}}
[data-testid="stMarkdown"] table {{
    width: 100%;
    border-collapse: collapse;
    font-size: 0.85rem;
    border-radius: 10px;
    overflow: hidden;
    border: 1px solid var(--border);
}}
[data-testid="stMarkdown"] th {{
    background: var(--bg3) !important;
    color: var(--accent2) !important;
    font-weight: 700 !important;
    padding: 10px 14px !important;
    text-align: left !important;
    font-size: 0.78rem !important;
    letter-spacing: 0.5px !important;
    text-transform: uppercase !important;
    border-bottom: 1px solid var(--border) !important;
}}
[data-testid="stMarkdown"] td {{
    padding: 10px 14px !important;
    border-bottom: 1px solid var(--border) !important;
    color: var(--text) !important;
    background: var(--bg) !important;
}}
[data-testid="stMarkdown"] tr:hover td {{
    background: var(--bg2) !important;
}}

[data-testid="stFileUploader"] {{
    margin-top: 10px !important;
}}

[data-testid="stAlert"] {{
    border-radius: 10px !important;
    border: none !important;
    font-size: 0.88rem !important;
}}

[data-testid="stAlert"] {{
    border-radius:10px!important;
    border:none!important;
    font-size:0.88rem!important;
}}
[data-testid="stAlert"] p {{
    color: var(--text) !important;
    font-weight: 600 !important;
    opacity: 1 !important;
}}
[data-testid="stAlert"] [data-testid="stMarkdownContainer"] p {{
    color: var(--text) !important;
}}
[data-testid="stTextInput"] {{
    margin-bottom: 0 !important;
    padding-bottom: 0 !important;
}}
[data-testid="stFileUploader"] {{
    margin-top: 0 !important;
}}
.left-panel > div > div > div {{
    gap: 12px !important;
}}

[data-testid="stDownloadButton"] button {{
    background: transparent !important;
    color: var(--accent) !important;
    border: 1.5px solid var(--accent) !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    padding: 8px 20px !important;
    margin-top: 0px !important;
    transition: all 0.2s !important;
}}
[data-testid="stDownloadButton"] button:hover {{
    background: var(--btn) !important;
    color: white !important;
}}

[data-testid="stDownloadButton"] {{
    margin-top: 22px !important;
}}

[data-testid="stSpinner"] {{ color: var(--accent) !important; }}

.success-banner {{
    background: var(--success);
    color: var(--success-t);
    border-radius: 8px;
    padding: 10px 16px;
    font-weight: 600;
    font-size: 0.88rem;
    margin-bottom: 22px;
    border: 1px solid var(--accent);
}}
.output-placeholder {{
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    height: 300px;
    color: var(--text2);
    font-size: 0.9rem;
    opacity: 0.5;
}}
.output-placeholder .icon {{ font-size: 3rem; margin-bottom: 12px; }}

.footer {{
    text-align: center;
    padding: 16px;
    color: var(--text2);
    font-size: 0.75rem;
    border-top: 1px solid var(--border);
    background: var(--bg2);
}}

.stImage img {{
    border-radius: 8px !important;
    border: 1px solid var(--border) !important;
}}

html, body, [data-testid="stAppViewContainer"] {{
    overflow-y: auto !important;
    overflow-x: hidden !important;
}}

::-webkit-scrollbar {{
    width: 0px;
    background: transparent;
}}
::-webkit-scrollbar-track {{ background: var(--bg); }}
::-webkit-scrollbar-thumb {{ background: var(--border); border-radius: 3px; }}
::-webkit-scrollbar-thumb:hover {{ background: var(--accent); }}

#MainMenu, footer, header {{ visibility: hidden; }}
</style>
""", unsafe_allow_html=True)

mode_icon = "☀️" if st.session_state.dark_mode else "🌙"
badges_html = "".join([
    f'<span class="badge">{b}</span>'
    for b in ["OpenRouter AI", "Vision AI", "Streamlit", "python-docx", "PyMuPDF"]
])

col_logo, col_toggle = st.columns([10, 1])
with col_logo:
    st.markdown(f"""
    <div class="navbar">
        <div class="nav-logo">📋 note2<span>action</span></div>
        <div class="nav-badges">{badges_html}</div>
    </div>
    """, unsafe_allow_html=True)
with col_toggle:
    if st.button(mode_icon, help="Toggle dark/light mode"):
        st.session_state.dark_mode = not st.session_state.dark_mode
        st.rerun()

SYSTEM_PROMPT = """
You are an expert meeting minutes generator. Analyze the provided meeting notes and produce a structured Minutes of Meeting (MoM).

OUTPUT FORMAT — Return a clean markdown table with these EXACT columns:
| Work Area | Action Item | Owner | Deadline | Priority | Risk / Blocker | Status |

RULES:
1. Group rows by Work Area (e.g., Construction, Finance, Design, IT, HR, General)
2. Every action item must be specific and clear
3. If deadline is not mentioned → write TBD
4. If owner is not mentioned → write Unassigned
5. Priority must be: High / Medium / Low
6. Risk/Blocker: any dependency or blocker. If none → write None
7. Status: default to Open
8. Extract ALL action items — do not miss any
9. After the table add:
## Summary
- Total action items found: X
- Owners assigned: X out of X
- High priority items: X

Return ONLY the markdown table and summary. No extra explanation.
"""

def extract_from_image(file):
    image = Image.open(file)
    return [image], "image"

def extract_from_pdf(file):
    pdf_bytes = file.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    images = []
    for page in doc:
        pix = page.get_pixmap(dpi=150)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        images.append(img)
    return images, "image"

def extract_from_docx(file):
    doc = Document(file)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return text, "text"

def extract_from_txt(file):
    text = file.read().decode("utf-8", errors="ignore")
    return text, "text"

def pil_to_base64(img):
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")

def call_ai(content, content_type):
    if content_type == "image":
        image_parts = []
        for img in content:
            b64 = pil_to_base64(img)
            image_parts.append({
                "type": "image_url",
                "image_url": {"url": f"data:image/png;base64,{b64}"}
            })
        messages = [{"role": "user", "content": [
            {"type": "text", "text": SYSTEM_PROMPT}, *image_parts
        ]}]
    else:
        messages = [{"role": "user", "content": f"{SYSTEM_PROMPT}\n\nMEETING NOTES:\n{content}"}]

    response = client.chat.completions.create(
        model="google/gemini-2.0-flash-lite-001",
        messages=messages
    )
    return response.choices[0].message.content

def parse_metrics(mom_text):
    rows = [l for l in mom_text.split("\n") if l.strip().startswith("|") and "---" not in l]
    action_items = max(len(rows) - 1, 0)
    owners = len(re.findall(r'\|\s*(?!Unassigned)(?!Owner)\s*[A-Z][a-zA-Z\s\.]+\s*\|', mom_text))
    high_priority = mom_text.lower().count("| high")
    time_saved = max(action_items * 2, 10)
    return {"action_items": action_items, "owners": owners,
            "high_priority": high_priority, "time_saved": time_saved}

def generate_docx(mom_text, meeting_name):
    doc = Document()
    doc.add_heading(f"Minutes of Meeting — {meeting_name}", 0)
    lines = mom_text.split("\n")
    table_lines = [l for l in lines if l.strip().startswith("|")]
    other_lines = [l for l in lines if not l.strip().startswith("|")]
    if table_lines:
        headers = [h.strip() for h in table_lines[0].split("|") if h.strip()]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for row_line in table_lines[2:]:
            cells = [c.strip() for c in row_line.split("|") if c.strip() != ""]
            if cells and len(cells) == len(headers):
                row = table.add_row().cells
                for i, cell in enumerate(cells):
                    row[i].text = cell
    doc.add_paragraph("")
    for line in other_lines:
        if line.strip():
            doc.add_paragraph(line.strip())
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

left, right = st.columns([4, 6], gap="small")
st.markdown('<style>div[data-testid="column"]{padding:0!important;}</style>', unsafe_allow_html=True)

with left:
    st.markdown('<div style="padding: 18px 28px 32px 28px;">', unsafe_allow_html=True)
    
    st.markdown('<div class="panel-title">Upload Meeting Notes</div>', unsafe_allow_html=True)

    meeting_name = st.text_input("Meeting name", placeholder="e.g. Q2 Review · Site Visit · Sprint Planning")
    st.markdown("<div style='margin-top:2px'></div>", unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "Drop your file here",
        type=["png", "jpg", "jpeg", "webp", "pdf", "docx", "txt"]
    )

    if uploaded_file:
        ext = uploaded_file.name.split(".")[-1].lower()
        st.success(f"✅ **{uploaded_file.name}** ({ext.upper()}) ready")
        if ext in ["png", "jpg", "jpeg", "webp"]:
            st.image(uploaded_file, caption="Preview", width=340)
        elif ext == "txt":
            preview = uploaded_file.read().decode("utf-8", errors="ignore")[:400]
            uploaded_file.seek(0)
            st.text_area("Preview", preview, height=130)
        elif ext == "docx":
            doc_prev = Document(uploaded_file)
            uploaded_file.seek(0)
            preview = "\n".join([p.text for p in doc_prev.paragraphs[:6] if p.text.strip()])
            st.text_area("Preview", preview, height=130)

    st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)
    generate_btn = st.button("⚡  Generate MoM", type="primary", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

with right:
    st.markdown('<div style="padding: 18px 36px 32px 36px;">', unsafe_allow_html=True)
    st.markdown('<div class="panel-title">Generated Output</div>', unsafe_allow_html=True)

    if generate_btn:
        if not uploaded_file:
            st.warning("⚠️ Please upload a file first.")
        elif not meeting_name.strip():
            st.warning("⚠️ Please enter a meeting name.")
        else:
            with st.spinner("🤖 AI is reading your notes..."):
                try:
                    ext = uploaded_file.name.split(".")[-1].lower()
                    if ext in ["png", "jpg", "jpeg", "webp"]:
                        content, ctype = extract_from_image(uploaded_file)
                    elif ext == "pdf":
                        content, ctype = extract_from_pdf(uploaded_file)
                    elif ext == "docx":
                        content, ctype = extract_from_docx(uploaded_file)
                    elif ext == "txt":
                        content, ctype = extract_from_txt(uploaded_file)
                    else:
                        st.error("Unsupported file type.")
                        st.stop()

                    mom_text = call_ai(content, ctype)
                    st.session_state["mom_text"] = mom_text
                    st.session_state["meeting_name"] = meeting_name
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")

    if "mom_text" in st.session_state:
        mom_text = st.session_state["mom_text"]
        meeting_name_saved = st.session_state["meeting_name"]
        
        st.markdown("""
        <style>
        ::-webkit-scrollbar {
            width: 6px !important;
        }
        ::-webkit-scrollbar-track {
            background: var(--bg);
        }
        ::-webkit-scrollbar-thumb {
            background: var(--border);
            border-radius: 3px;
        }
        </style>
        """, unsafe_allow_html=True)

        st.markdown('<div class="success-banner">✅ MoM Generated successfully!</div>', unsafe_allow_html=True)

        metrics = parse_metrics(mom_text)
        st.markdown(f"""
        <div class="metrics-row">
            <div class="metric-card">
                <div class="metric-value">{metrics['action_items']}</div>
                <div class="metric-label">Action Items</div>
            </div>
            <div class="metric-card">
                <div class="metric-value">{metrics['owners']}</div>
                <div class="metric-label">Owners Assigned</div>
            </div>
            <div class="metric-card">
                <div class="metric-value">{metrics['high_priority']}</div>
                <div class="metric-label">High Priority</div>
            </div>
            <div class="metric-card">
                <div class="metric-value">~{metrics['time_saved']}m</div>
                <div class="metric-label">Time Saved</div>
            </div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown(mom_text)

        docx_buf = generate_docx(mom_text, meeting_name_saved)
        st.download_button(
            label="⬇️  Download as Word (.docx)",
            data=docx_buf,
            file_name=f"{meeting_name_saved.replace(' ', '_')}_MoM.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.markdown("""
        <div class="output-placeholder">
            <div class="icon">📄</div>
            Upload your meeting notes and click Generate MoM
        </div>
        """, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

st.markdown("""
<div class="footer">
    Built with OpenRouter AI · Vision AI · Streamlit · © 2026 · Vishwa Desai
    <strong>note2action</strong> — turning chaos into clarity
</div>
""", unsafe_allow_html=True)
